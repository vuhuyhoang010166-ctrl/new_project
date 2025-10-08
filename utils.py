# -*- coding: utf-8 -*-
"""
Module chứa các hàm tiện ích
"""

import io
import streamlit as st
from docx import Document
from typing import Optional
from config import ERROR_MESSAGES


class DocumentReader:
    """Class để đọc các loại tài liệu"""

    @staticmethod
    def extract_text_from_docx(uploaded_file) -> Optional[str]:
        """
        Đọc và trích xuất toàn bộ văn bản từ file .docx

        Args:
            uploaded_file: UploadedFile object từ Streamlit

        Returns:
            String chứa nội dung văn bản hoặc None nếu có lỗi
        """
        try:
            # Đọc file vào BytesIO để không cần lưu xuống đĩa
            file_bytes = uploaded_file.read()
            document = Document(io.BytesIO(file_bytes))

            # Trích xuất text từ các paragraphs
            full_text = [para.text for para in document.paragraphs if para.text.strip()]

            # Trích xuất text từ tables (nếu có)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            # Reset file pointer để có thể đọc lại nếu cần
            uploaded_file.seek(0)

            text_content = '\n'.join(full_text)

            if not text_content.strip():
                return None

            return text_content

        except Exception as e:
            st.error(ERROR_MESSAGES["file_read_error"].format(str(e)))
            return None


class SessionStateManager:
    """Class quản lý session state của Streamlit"""

    @staticmethod
    def initialize_session_state():
        """Khởi tạo tất cả các biến session state cần thiết"""
        if 'project_data' not in st.session_state:
            st.session_state.project_data = None

        if 'cash_flow_df' not in st.session_state:
            st.session_state.cash_flow_df = None

        if 'metrics' not in st.session_state:
            st.session_state.metrics = None

        if 'analysis_requested' not in st.session_state:
            st.session_state.analysis_requested = False

        if 'ai_analysis_result' not in st.session_state:
            st.session_state.ai_analysis_result = None

        if 'uploaded_file_content' not in st.session_state:
            st.session_state.uploaded_file_content = None

        if 'uploaded_file_name' not in st.session_state:
            st.session_state.uploaded_file_name = None

    @staticmethod
    def reset_calculation_state():
        """Reset các kết quả tính toán khi có dữ liệu mới"""
        st.session_state.cash_flow_df = None
        st.session_state.metrics = None
        st.session_state.analysis_requested = False
        st.session_state.ai_analysis_result = None

    @staticmethod
    def reset_all_state():
        """Reset toàn bộ session state"""
        st.session_state.project_data = None
        st.session_state.cash_flow_df = None
        st.session_state.metrics = None
        st.session_state.analysis_requested = False
        st.session_state.ai_analysis_result = None
        st.session_state.uploaded_file_content = None
        st.session_state.uploaded_file_name = None


class DataFormatter:
    """Class format dữ liệu để hiển thị"""

    @staticmethod
    def format_currency(value: float) -> str:
        """Format số tiền theo định dạng Việt Nam"""
        return f"{value:,.0f} VNĐ"

    @staticmethod
    def format_percentage(value: float) -> str:
        """Format phần trăm"""
        return f"{value:.2f}%"

    @staticmethod
    def format_year(value: float) -> str:
        """Format số năm"""
        return f"{value:.2f} năm"

    @staticmethod
    def format_metric_value(value, metric_type: str) -> str:
        """
        Format giá trị metric theo loại

        Args:
            value: Giá trị cần format
            metric_type: Loại metric ('NPV', 'IRR', 'PP', 'DPP')

        Returns:
            String đã được format
        """
        if isinstance(value, str):
            return value

        if metric_type == 'NPV':
            return DataFormatter.format_currency(value)
        elif metric_type == 'IRR':
            return DataFormatter.format_percentage(value)
        elif metric_type in ['PP', 'DPP']:
            return DataFormatter.format_year(value)
        else:
            return str(value)


def get_api_key_from_secrets_or_input() -> Optional[str]:
    """
    Lấy API key từ Streamlit secrets hoặc input của người dùng

    Returns:
        API key string hoặc None
    """
    try:
        # Thử lấy từ secrets trước
        api_key = st.secrets.get("GEMINI_API_KEY", None)
        if api_key:
            st.success("✅ Đã tìm thấy API Key từ cấu hình.", icon="🔑")
            return api_key
    except:
        pass

    # Nếu không có trong secrets, yêu cầu nhập
    api_key = st.text_input(
        "Nhập Gemini API Key của bạn:",
        type="password",
        help="Lấy API key tại: https://makersuite.google.com/app/apikey"
    )

    return api_key if api_key else None
